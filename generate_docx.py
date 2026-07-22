"""
ProjectPilot AI — Markdown to DOCX Generator Script
--------------------------------------------------
This script reads PROJECT_EXPLANATION.md and generates a styled Microsoft Word (.docx) document.
Requirements: pip install python-docx
"""

import os
import sys

def convert_md_to_docx():
    try:
        import docx
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except ImportError:
        print("[-] The 'python-docx' library is not installed.")
        print("[+] Please run: pip install python-docx")
        return

    md_path = os.path.join(os.path.dirname(__file__), "PROJECT_EXPLANATION.md")
    output_docx = os.path.join(os.path.dirname(__file__), "ProjectPilot_AI_Complete_Documentation.docx")

    if not os.path.exists(md_path):
        print(f"[-] Markdown file not found at: {md_path}")
        return

    doc = docx.Document()

    # Page Margins (1 inch)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Style colors
    PRIMARY_COLOR = RGBColor(79, 70, 229)    # #4f46e5 Indigo
    SECONDARY_COLOR = RGBColor(124, 58, 237) # #7c3aed Purple
    TEXT_DARK = RGBColor(15, 23, 42)        # #0f172a Slate 900
    TEXT_MUTED = RGBColor(51, 65, 85)       # #334155 Slate 700

    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    in_code_block = False
    code_buffer = []

    for line in lines:
        stripped = line.strip()

        # Handle Code Blocks
        if stripped.startswith("```"):
            if in_code_block:
                # End code block
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.4)
                run = p.add_run("\n".join(code_buffer))
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(30, 41, 59)
                code_buffer = []
                in_code_block = False
            else:
                in_code_block = True
            continue

        if in_code_block:
            code_buffer.append(line.rstrip())
            continue

        # Handle Headings
        if stripped.startswith("# "):
            p = doc.add_heading(stripped[2:], level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(22)
                run.font.bold = True
                run.font.color.rgb = PRIMARY_COLOR
            doc.add_paragraph() # Spacer
        elif stripped.startswith("## "):
            p = doc.add_heading(stripped[3:], level=2)
            for run in p.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(16)
                run.font.bold = True
                run.font.color.rgb = SECONDARY_COLOR
        elif stripped.startswith("### "):
            p = doc.add_heading(stripped[4:], level=3)
            for run in p.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(13)
                run.font.bold = True
                run.font.color.rgb = TEXT_DARK
        elif stripped.startswith("#### "):
            p = doc.add_heading(stripped[5:], level=4)
            for run in p.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(11)
                run.font.bold = True
                run.font.color.rgb = TEXT_MUTED
        elif stripped.startswith("> "):
            # Blockquote
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            run = p.add_run(stripped[2:])
            run.font.italic = True
            run.font.size = Pt(10)
            run.font.color.rgb = TEXT_MUTED
        elif stripped.startswith("* ") or stripped.startswith("- "):
            # Bullet list
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(stripped[2:])
            run.font.name = 'Arial'
            run.font.size = Pt(10.5)
            run.font.color.rgb = TEXT_DARK
        elif stripped == "---":
            # Horizontal divider
            p = doc.add_paragraph("_________________________________________________________________________________")
            for run in p.runs:
                run.font.color.rgb = RGBColor(226, 232, 240)
        elif stripped:
            # Normal Paragraph
            p = doc.add_paragraph()
            run = p.add_run(stripped)
            run.font.name = 'Arial'
            run.font.size = Pt(10.5)
            run.font.color.rgb = TEXT_DARK

    doc.save(output_docx)
    print(f"[+] Successfully generated DOCX report at: {output_docx}")

if __name__ == "__main__":
    convert_md_to_docx()
